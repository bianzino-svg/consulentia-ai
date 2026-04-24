from fastapi import FastAPI, Query
from fastapi.responses import HTMLResponse, JSONResponse, FileResponse
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Inches
import yfinance as yf
import urllib.request
import xml.etree.ElementTree as ET
import urllib.parse
import json
import re
import os

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image, PageBreak
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

app = FastAPI()

BASE_DIR = Path(__file__).resolve().parent
ARCHIVE_DIR = BASE_DIR / "archivio_report"
ARCHIVE_DIR.mkdir(exist_ok=True)

LOGO_PATH = BASE_DIR / "logo_amb.png"


PROFILE_UNIVERSES = {
    "prudente": {
        "label": "Prudente",
        "equity_ticker": "VWCE.DE",
        "style": "Approccio orientato alla protezione del capitale, con crescita graduale e volatilità contenuta.",
        "base_allocation": {
            "Azionario": 25,
            "Obbligazionario": 55,
            "Commodities": 10,
            "Liquidità": 10,
        },
        "tactical_allocations": {
            "risk_on": {"Azionario": 30, "Obbligazionario": 52, "Commodities": 10, "Liquidità": 8},
            "neutral": {"Azionario": 25, "Obbligazionario": 55, "Commodities": 10, "Liquidità": 10},
            "risk_off": {"Azionario": 20, "Obbligazionario": 58, "Commodities": 10, "Liquidità": 12},
        },
        "equity_funds": [
            "Fidelity Funds - European Dynamic Growth Fund",
            "Pictet - Global Megatrend Selection",
            "JPMorgan Funds - Global Focus Fund",
            "Morgan Stanley Investment Funds - Global Opportunity Fund",
        ],
        "bond_funds": [
            "Franklin Templeton Investment Funds - Global Bond Fund",
            "Nordea 1 - European Covered Bond Fund",
            "Carmignac Sécurité",
            "Amundi Funds - Global Aggregate Bond",
        ],
        "commodity_funds": [
            "BlackRock Global Funds - World Gold Fund",
            "Schroder ISF - Global Energy",
        ],
        "abs_return_funds": [],
    },
    "bilanciato": {
        "label": "Bilanciato",
        "equity_ticker": "SWDA.L",
        "style": "Approccio equilibrato: ricerca di rendimento con attenzione alla tenuta complessiva del portafoglio.",
        "base_allocation": {
            "Azionario": 45,
            "Obbligazionario": 35,
            "Commodities": 10,
            "Liquidità": 10,
        },
        "tactical_allocations": {
            "risk_on": {"Azionario": 50, "Obbligazionario": 30, "Commodities": 10, "Liquidità": 10},
            "neutral": {"Azionario": 45, "Obbligazionario": 35, "Commodities": 10, "Liquidità": 10},
            "risk_off": {"Azionario": 38, "Obbligazionario": 40, "Commodities": 10, "Liquidità": 12},
        },
        "equity_funds": [
            "Fidelity Funds - European Dynamic Growth Fund",
            "Pictet - Global Megatrend Selection",
            "MS INVF - Global Opportunity Fund",
            "JPMorgan Funds - Global Focus Fund",
        ],
        "bond_funds": [
            "Franklin Templeton Investment Funds - Global Bond Fund",
            "Nordea 1 - European Covered Bond Fund",
            "Carmignac Sécurité",
            "Amundi Funds - Global Aggregate Bond",
        ],
        "commodity_funds": [
            "BlackRock Global Funds - World Gold Fund",
            "Schroder ISF - Global Energy",
        ],
        "abs_return_funds": [],
    },
    "dinamico": {
        "label": "Dinamico",
        "equity_ticker": "SWDA.L",
        "style": "Approccio più deciso: maggiore esposizione al rischio, ma con controllo della diversificazione.",
        "base_allocation": {
            "Azionario": 60,
            "Obbligazionario": 25,
            "Commodities": 10,
            "Absolute / Total Return": 5,
        },
        "tactical_allocations": {
            "risk_on": {"Azionario": 65, "Obbligazionario": 20, "Commodities": 10, "Absolute / Total Return": 5},
            "neutral": {"Azionario": 60, "Obbligazionario": 25, "Commodities": 10, "Absolute / Total Return": 5},
            "risk_off": {"Azionario": 52, "Obbligazionario": 30, "Commodities": 13, "Absolute / Total Return": 5},
        },
        "equity_funds": [
            "Morgan Stanley Investment Funds - Global Opportunity Fund",
            "Pictet - Global Megatrend Selection",
            "Fidelity Funds - European Dynamic Growth Fund",
            "JPMorgan Funds - Global Focus Fund",
        ],
        "bond_funds": [
            "Franklin Templeton Investment Funds - Global Bond Fund",
            "Nordea 1 - European Covered Bond Fund",
            "Amundi Funds - Global Aggregate Bond",
            "Carmignac Sécurité",
        ],
        "commodity_funds": [
            "BlackRock Global Funds - World Gold Fund",
            "Schroder ISF - Global Energy",
        ],
        "abs_return_funds": [
            "DNCA Invest - Alpha Bonds",
        ],
    },
}

DISCLAIMER_TEXT = (
    "Disclaimer: il presente report è frutto di ricerca ed elaborazione tramite software proprietario. "
    "Ha finalità esclusivamente informative, di supporto all'analisi e di discussione professionale. "
    "Non costituisce consulenza personalizzata, raccomandazione individuale, sollecitazione all'investimento, "
    "offerta o invito all'acquisto o alla vendita di strumenti finanziari. "
    "Le analisi quantitative, tecniche, macroeconomiche e previsionali hanno natura indicativa e non garantiscono risultati futuri. "
    "I mercati possono evolvere in modo differente rispetto agli scenari ipotizzati. "
    "Prima di assumere decisioni operative è sempre opportuno verificare l'adeguatezza dello strumento, "
    "il profilo di rischio, i costi, la documentazione ufficiale e l'orizzonte temporale dell'investitore."
)


def safe_float(value):
    try:
        return float(value)
    except Exception:
        return None


def list_lines(items):
    return "\n".join([f"- {x}" for x in items]) if items else "- Nessuno"


def fetch_json(url, timeout=15):
    with urllib.request.urlopen(url, timeout=timeout) as response:
        return json.loads(response.read().decode("utf-8"))


def weighted_projection(values):
    clean = [safe_float(x) for x in values if safe_float(x) is not None]
    if not clean:
        return None, None
    if len(clean) == 1:
        return clean[-1], clean[-1]

    n = len(clean)
    xs = list(range(n))
    weights = [i + 1 for i in range(n)]

    w_sum = sum(weights)
    x_mean = sum(x * w for x, w in zip(xs, weights)) / w_sum
    y_mean = sum(y * w for y, w in zip(clean, weights)) / w_sum

    denom = sum(w * ((x - x_mean) ** 2) for x, w in zip(xs, weights))
    if denom == 0:
        return clean[-1], clean[-1]

    slope = sum(w * (x - x_mean) * (y - y_mean) for x, y, w in zip(xs, clean, weights)) / denom
    projection = clean[-1] + slope
    projection = (projection * 0.7) + (clean[-1] * 0.3)

    return clean[-1], projection


def format_pair(current, projection, suffix=""):
    c = "n.d." if current is None else f"{round(current, 2)}{suffix}"
    p = "n.d." if projection is None else f"{round(projection, 2)}{suffix}"
    return f"{c} → {p}"


def clean_text_for_pdf(text: str) -> str:
    replacements = {
        "🟢": "[VERDE]",
        "🟡": "[GIALLO]",
        "🔴": "[ROSSO]",
        "⚪": "[ND]",
        "→": "->",
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return re.sub(r"[^\x00-\x7FàèéìòùÀÈÉÌÒÙ€*%/()\-.,:; ]", "", text)


def get_history(ticker_symbol, period="3mo"):
    try:
        ticker = yf.Ticker(ticker_symbol)
        hist = ticker.history(period=period)
        if hist.empty:
            return None
        return hist
    except Exception:
        return None


def get_last_price_and_trend(ticker_symbol, name):
    hist = get_history(ticker_symbol, period="5d")
    if hist is None or "Close" not in hist.columns or len(hist) < 2:
        return {"name": name, "last": "n.d.", "trend": "n.d."}

    first_val = float(hist["Close"].iloc[0])
    last_val = float(hist["Close"].iloc[-1])
    trend = "positivo" if last_val > first_val else "negativo"

    return {"name": name, "last": round(last_val, 2), "trend": trend}


def get_top_news():
    feeds = [
        "https://feeds.reuters.com/reuters/businessNews",
        "https://feeds.reuters.com/reuters/worldNews",
    ]

    news_items = []

    for feed_url in feeds:
        try:
            with urllib.request.urlopen(feed_url, timeout=8) as response:
                xml_data = response.read()

            root = ET.fromstring(xml_data)

            for item in root.findall(".//item")[:4]:
                title_el = item.find("title")
                if title_el is not None and title_el.text:
                    title = title_el.text.strip()
                    if title and title not in news_items:
                        news_items.append(title)
                if len(news_items) >= 6:
                    break
        except Exception:
            continue

        if len(news_items) >= 6:
            break

    if not news_items:
        return ["Nessuna news disponibile al momento."]

    return news_items[:6]


def build_market_outlook(sp_trend, nasdaq_trend, gold_trend, oil_trend):
    if sp_trend == "positivo" and nasdaq_trend == "positivo":
        outlook = "Il quadro generale resta costruttivo: gli indici principali tengono bene e il tono di mercato rimane favorevole."
        market_mode = "risk_on"
    elif sp_trend == "negativo" and nasdaq_trend == "negativo":
        outlook = "Il mercato è più fragile del solito: conviene tenere un passo prudente e non forzare nuovi ingressi."
        market_mode = "risk_off"
    else:
        outlook = "Siamo in una fase mista: il mercato manda segnali contrastanti e vale la pena muoversi con selettività."
        market_mode = "neutral"

    extra_notes = []
    if gold_trend == "positivo":
        extra_notes.append("L'oro sta lavorando bene come copertura e continua a meritare attenzione.")
    if oil_trend == "positivo":
        extra_notes.append("Il petrolio in rialzo può mantenere viva la sensibilità sul tema inflazione.")
    if not extra_notes:
        extra_notes.append("Sul fronte intermarket non emerge un segnale dominante particolarmente forte.")

    return outlook, market_mode, extra_notes


def find_last_fractals(hist):
    if hist is None or len(hist) < 7:
        return None, None

    highs = hist["High"].tolist()
    lows = hist["Low"].tolist()
    dates = hist.index.tolist()

    last_up = None
    last_down = None

    for i in range(2, len(hist) - 2):
        high = highs[i]
        low = lows[i]

        is_up = (
            high > highs[i - 1] and
            high > highs[i - 2] and
            high > highs[i + 1] and
            high > highs[i + 2]
        )

        is_down = (
            low < lows[i - 1] and
            low < lows[i - 2] and
            low < lows[i + 1] and
            low < lows[i + 2]
        )

        if is_up:
            last_up = {
                "date": dates[i].strftime("%Y-%m-%d"),
                "level": round(float(high), 2)
            }

        if is_down:
            last_down = {
                "date": dates[i].strftime("%Y-%m-%d"),
                "level": round(float(low), 2)
            }

    return last_up, last_down


def get_traffic_light(price, sma20, up_level, down_level):
    if isinstance(price, str):
        return "⚪", "dati non disponibili"

    if up_level is not None and price > up_level and price > sma20:
        return "🟢", "setup favorevole"
    if down_level is not None and price < down_level and price < sma20:
        return "🔴", "debolezza / alleggerire"
    return "🟡", "attesa / conferma"


def analyze_fractal_signal(ticker_symbol, label):
    hist = get_history(ticker_symbol, period="3mo")
    if hist is None or len(hist) < 25:
        return {
            "label": label,
            "status": "dati insufficienti",
            "price": "n.d.",
            "up_fractal": "n.d.",
            "down_fractal": "n.d.",
            "light": "⚪",
            "light_text": "dati insufficienti",
            "comment": "Storico insufficiente per una lettura tecnica affidabile."
        }

    close = float(hist["Close"].iloc[-1])
    sma20 = float(hist["Close"].tail(20).mean())

    last_up, last_down = find_last_fractals(hist)

    if last_up is None and last_down is None:
        return {
            "label": label,
            "status": "nessun frattale valido",
            "price": round(close, 2),
            "up_fractal": "n.d.",
            "down_fractal": "n.d.",
            "light": "⚪",
            "light_text": "nessun pattern valido",
            "comment": "Nel periodo osservato non si è formato un pattern frattale davvero utile."
        }

    up_level = last_up["level"] if last_up else None
    down_level = last_down["level"] if last_down else None

    if up_level is not None and close > up_level and close > sma20:
        status = "rialzista"
        comment = "Prezzo sopra l'ultimo frattale rialzista e sopra media 20 giorni."
    elif down_level is not None and close < down_level and close < sma20:
        status = "ribassista"
        comment = "Prezzo sotto l'ultimo frattale ribassista e sotto media 20 giorni."
    else:
        status = "neutrale"
        comment = "Prezzo ancora in mezzo ai livelli chiave: serve conferma."

    light, light_text = get_traffic_light(close, sma20, up_level, down_level)

    return {
        "label": label,
        "status": status,
        "price": round(close, 2),
        "up_fractal": f"{last_up['level']} ({last_up['date']})" if last_up else "n.d.",
        "down_fractal": f"{last_down['level']} ({last_down['date']})" if last_down else "n.d.",
        "light": light,
        "light_text": light_text,
        "comment": comment
    }


def build_fractal_section():
    analyses = [
        analyze_fractal_signal("^GSPC", "S&P500"),
        analyze_fractal_signal("^IXIC", "Nasdaq"),
        analyze_fractal_signal("GC=F", "Oro"),
        analyze_fractal_signal("EURUSD=X", "EUR/USD"),
        analyze_fractal_signal("BTC-USD", "Bitcoin"),
        analyze_fractal_signal("ETH-USD", "Ethereum"),
    ]

    green = sum(1 for x in analyses if x["light"] == "🟢")
    red = sum(1 for x in analyses if x["light"] == "🔴")
    yellow = sum(1 for x in analyses if x["light"] == "🟡")

    if green >= 4:
        summary = "Quadro tecnico complessivamente favorevole."
    elif red >= 4:
        summary = "Quadro tecnico complessivamente debole."
    else:
        summary = "Quadro tecnico misto: meglio aspettare conferme dove serve."

    traffic_summary = f"Semafori: Verde {green} | Giallo {yellow} | Rosso {red}"
    return analyses, summary, traffic_summary


def build_operational_actions(market_mode, analyses):
    actions = []

    sp_signal = next((a for a in analyses if a["label"] == "S&P500"), None)
    nasdaq_signal = next((a for a in analyses if a["label"] == "Nasdaq"), None)
    gold_signal = next((a for a in analyses if a["label"] == "Oro"), None)
    eurusd_signal = next((a for a in analyses if a["label"] == "EUR/USD"), None)
    btc_signal = next((a for a in analyses if a["label"] == "Bitcoin"), None)
    eth_signal = next((a for a in analyses if a["label"] == "Ethereum"), None)

    if market_mode == "risk_on":
        actions.append("Si può aumentare il rischio in modo graduale, senza entrare tutto insieme.")
    elif market_mode == "risk_off":
        actions.append("Qui conviene alleggerire un po' il profilo di rischio e privilegiare tenuta e qualità.")
    else:
        actions.append("Meglio muoversi con calma e costruire a piccoli passi, senza forzare il timing.")

    green_count = sum(1 for a in analyses if a["light"] == "🟢")
    red_count = sum(1 for a in analyses if a["light"] == "🔴")

    if green_count >= 4:
        actions.append("Il quadro tecnico aggregato consente un atteggiamento costruttivo, ma sempre graduale.")
    elif red_count >= 4:
        actions.append("Il quadro tecnico aggregato suggerisce un taglio più prudente e un controllo stretto del rischio.")
    else:
        actions.append("Il quadro tecnico aggregato è misto: meglio evitare sovraesposizioni e restare selettivi.")

    if gold_signal:
        if gold_signal["light"] == "🟢":
            actions.append("L'oro continua a fare bene il suo mestiere di copertura.")
        elif gold_signal["light"] == "🔴":
            actions.append("L'oro nel breve è meno forte, ma resta utile come diversificazione strategica.")

    if sp_signal and nasdaq_signal:
        if sp_signal["light"] == "🟢" and nasdaq_signal["light"] == "🟢":
            actions.append("Azionario USA ben allineato: base costruttiva.")
        elif sp_signal["light"] == "🔴" and nasdaq_signal["light"] == "🔴":
            actions.append("Azionario USA debole: meglio una postura più difensiva.")
        else:
            actions.append("Azionario USA non perfettamente allineato: serve conferma.")

    if eurusd_signal:
        if eurusd_signal["light"] == "🟢":
            actions.append("EUR/USD in miglioramento tecnico: utile monitorare l'effetto cambio.")
        elif eurusd_signal["light"] == "🔴":
            actions.append("EUR/USD in debolezza tecnica: attenzione alla componente valutaria.")

    if btc_signal or eth_signal:
        crypto_green = sum(1 for x in [btc_signal, eth_signal] if x and x["light"] == "🟢")
        crypto_red = sum(1 for x in [btc_signal, eth_signal] if x and x["light"] == "🔴")

        if crypto_green == 2:
            actions.append("Bitcoin ed Ethereum sono entrambi ben impostati, ma il comparto resta da trattare con prudenza.")
        elif crypto_red == 2:
            actions.append("Bitcoin ed Ethereum mostrano debolezza tecnica diffusa.")
        else:
            actions.append("Il comparto crypto manda segnali misti: meglio disciplina e selettività.")

    return actions


WB_COUNTRIES = {
    "USA": ("USA", "USA"),
    "CHN": ("CHN", "Cina"),
    "EMU": ("EMU", "Euro Area"),
    "DEU": ("DEU", "Germania"),
    "ITA": ("ITA", "Italia"),
}

WB_INDICATORS = {
    "inflation": "FP.CPI.TOTL.ZG",
    "gdp": "NY.GDP.MKTP.KD.ZG",
    "consumption": "NE.CON.PRVT.KD.ZG",
    "unemployment": "SL.UEM.TOTL.ZS",
    "jobs": "SL.EMP.TOTL.SP.ZS",
    "rates": "FR.INR.LEND",
}


def fetch_world_bank_series(country_code, indicator_code, per_page=80):
    try:
        url = (
            f"https://api.worldbank.org/v2/country/{country_code}/indicator/{indicator_code}"
            f"?format=json&per_page={per_page}"
        )
        payload = fetch_json(url)
        if not isinstance(payload, list) or len(payload) < 2:
            return []
        rows = payload[1]
        out = []
        for row in rows:
            year = row.get("date")
            value = row.get("value")
            if year is None or value is None:
                continue
            try:
                out.append((int(year), float(value)))
            except Exception:
                continue
        out.sort(key=lambda x: x[0])
        return out
    except Exception:
        return []


def fetch_ecb_rate():
    try:
        url = "https://data-api.ecb.europa.eu/service/data/FM/D.U2.EUR.4F.KR.DFR.LEV?format=jsondata&lastNObservations=1"
        payload = fetch_json(url)
        series = payload.get("dataSets", [{}])[0].get("series", {})
        if not series:
            return None
        first_series = next(iter(series.values()))
        obs = first_series.get("observations", {})
        if not obs:
            return None
        first_obs = next(iter(obs.values()))
        if isinstance(first_obs, list) and first_obs:
            return safe_float(first_obs[0])
        return None
    except Exception:
        return None


def build_macro_comment(row):
    notes = []

    inf = row.get("inflation_current_num")
    gdp = row.get("gdp_current_num")
    unemp = row.get("unemployment_current_num")

    if inf is not None:
        if inf > 4:
            notes.append("inflazione alta")
        elif inf < 2:
            notes.append("inflazione più sotto controllo")

    if gdp is not None:
        if gdp > 2:
            notes.append("crescita discreta")
        elif gdp < 0:
            notes.append("crescita debole")

    if unemp is not None:
        if unemp < 5:
            notes.append("mercato del lavoro solido")
        elif unemp > 8:
            notes.append("mercato del lavoro fragile")

    if not notes:
        return "quadro macro abbastanza neutro"

    return ", ".join(notes)


def synthetic_leading_from_series(gdp_current, gdp_proj, cons_current, cons_proj, un_current, un_proj):
    score = 100.0
    if gdp_current is not None and gdp_proj is not None:
        score += (gdp_proj - gdp_current) * 1.8
    if cons_current is not None and cons_proj is not None:
        score += (cons_proj - cons_current) * 1.2
    if un_current is not None and un_proj is not None:
        score -= (un_proj - un_current) * 1.5
    return round(score, 2), round(score + ((gdp_proj or 0) - (gdp_current or 0)) * 0.4, 2)


def build_macro_row(country_key):
    wb_code, label = WB_COUNTRIES[country_key]

    series_map = {}
    for k, code in WB_INDICATORS.items():
        series_map[k] = fetch_world_bank_series(wb_code, code)

    inf_current, inf_proj = weighted_projection([v for _, v in series_map["inflation"]][-10:])
    gdp_current, gdp_proj = weighted_projection([v for _, v in series_map["gdp"]][-10:])
    cons_current, cons_proj = weighted_projection([v for _, v in series_map["consumption"]][-10:])
    un_current, un_proj = weighted_projection([v for _, v in series_map["unemployment"]][-10:])
    jobs_current, jobs_proj = weighted_projection([v for _, v in series_map["jobs"]][-10:])

    if country_key == "EMU":
        rate_current = fetch_ecb_rate()
        rate_proj = rate_current
    elif country_key in ["USA", "CHN"]:
        rate_current, rate_proj = weighted_projection([v for _, v in series_map["rates"]][-10:])
    else:
        rate_current, rate_proj = None, None

    lead_current, lead_proj = synthetic_leading_from_series(
        gdp_current, gdp_proj, cons_current, cons_proj, un_current, un_proj
    )

    row = {
        "country": label,
        "inflation_current_num": inf_current,
        "inflation_projection_num": inf_proj,
        "gdp_current_num": gdp_current,
        "gdp_projection_num": gdp_proj,
        "rates_current_num": rate_current,
        "rates_projection_num": rate_proj,
        "consumption_current_num": cons_current,
        "consumption_projection_num": cons_proj,
        "unemployment_current_num": un_current,
        "unemployment_projection_num": un_proj,
        "jobs_current_num": jobs_current,
        "jobs_projection_num": jobs_proj,
        "leading_current_num": lead_current,
        "leading_projection_num": lead_proj,
    }

    row["inflation"] = format_pair(inf_current, inf_proj, "%")
    row["gdp"] = format_pair(gdp_current, gdp_proj, "%")
    row["rates"] = format_pair(rate_current, rate_proj, "%")
    row["consumption"] = format_pair(cons_current, cons_proj, "%")
    row["unemployment"] = format_pair(un_current, un_proj, "%")
    row["jobs"] = format_pair(jobs_current, jobs_proj, "%")
    row["leading"] = format_pair(lead_current, lead_proj, "")
    row["comment"] = build_macro_comment(row)

    return row


def build_macro_table():
    rows = []
    for key in ["USA", "CHN", "EMU", "DEU", "ITA"]:
        try:
            rows.append(build_macro_row(key))
        except Exception:
            rows.append({
                "country": WB_COUNTRIES[key][1],
                "inflation": "n.d.",
                "gdp": "n.d.",
                "rates": "n.d.",
                "consumption": "n.d.",
                "unemployment": "n.d.",
                "jobs": "n.d.",
                "leading": "n.d.",
                "comment": "dati momentaneamente non disponibili"
            })
    return rows


def snapshot_path_from_txt(txt_path: Path):
    return txt_path.with_suffix(".json")


def build_snapshot(profile: str, data: dict):
    fractal_lights = {}
    for item in data.get("fractal_analyses", []):
        fractal_lights[item["label"]] = item["light"]

    return {
        "profile": profile.lower().strip(),
        "timestamp": data.get("timestamp"),
        "base_allocation": data.get("base_allocation", {}),
        "final_allocation": data.get("final_allocation", {}),
        "equity_funds": data.get("equity_funds", []),
        "bond_funds": data.get("bond_funds", []),
        "commodity_funds": data.get("commodity_funds", []),
        "abs_return_funds": data.get("abs_return_funds", []),
        "outlook": data.get("outlook", ""),
        "fractal_lights": fractal_lights,
    }


def save_snapshot(profile: str, data: dict, txt_path: Path):
    snapshot = build_snapshot(profile, data)
    with open(snapshot_path_from_txt(txt_path), "w", encoding="utf-8") as f:
        json.dump(snapshot, f, ensure_ascii=False, indent=2)


def get_latest_snapshot(profile: str):
    profile = profile.lower().strip()
    candidates = []
    for path in ARCHIVE_DIR.rglob("*.json"):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if data.get("profile") == profile:
                candidates.append((path, data))
        except Exception:
            continue

    if not candidates:
        return None

    candidates.sort(key=lambda x: x[0].stat().st_mtime, reverse=True)
    return candidates[0][1]


def build_allocation_change_flags(current_alloc: dict, previous_alloc: dict | None):
    flags = {}
    if not previous_alloc:
        for k in current_alloc.keys():
            flags[k] = False
        return flags

    all_keys = set(current_alloc.keys()) | set(previous_alloc.keys())
    for key in all_keys:
        flags[key] = current_alloc.get(key) != previous_alloc.get(key)
    return flags


def build_list_change_flags(current_list, previous_list):
    current_set = set(current_list or [])
    previous_set = set(previous_list or [])
    return {
        "entered": sorted(list(current_set - previous_set)),
        "exited": sorted(list(previous_set - current_set)),
        "changed": current_set != previous_set
    }


def build_outlook_change_flag(current_outlook, previous_outlook):
    if not previous_outlook:
        return False
    return current_outlook != previous_outlook


def build_fractal_change_flags(current_analyses, previous_fractal_lights):
    flags = {}
    previous_fractal_lights = previous_fractal_lights or {}
    for item in current_analyses:
        old_light = previous_fractal_lights.get(item["label"])
        flags[item["label"]] = (old_light is not None and old_light != item["light"])
    return flags


def build_strategy_text(profile_data, market_mode):
    base_style = profile_data["style"]

    if market_mode == "risk_on":
        return (
            f"{base_style}\n"
            "Il contesto consente di essere un po' più costruttivi, ma sempre con ingressi graduali e senza inseguire il mercato.\n"
            "Meglio privilegiare strumenti liquidi, diversificati e facilmente monitorabili."
        )
    elif market_mode == "risk_off":
        return (
            f"{base_style}\n"
            "Qui il punto non è rincorrere rendimento, ma difendere bene il portafoglio e lasciare spazio di manovra.\n"
            "In questa fase conviene alzare la qualità media degli strumenti e tenere una gestione più disciplinata."
        )
    else:
        return (
            f"{base_style}\n"
            "Lo scenario è intermedio: non è il momento di strafare, ma nemmeno di restare completamente fermi.\n"
            "Ha senso costruire con calma, con attenzione alla qualità dei fondi e alla coerenza del profilo."
        )


def choose_profile_data(profile, market_mode):
    selected = PROFILE_UNIVERSES.get(profile, PROFILE_UNIVERSES["bilanciato"]).copy()
    selected["final_allocation"] = selected["tactical_allocations"][market_mode]
    return selected


def build_dashboard_data(profile: str):
    sp500 = get_last_price_and_trend("^GSPC", "S&P500")
    nasdaq = get_last_price_and_trend("^IXIC", "Nasdaq")
    gold = get_last_price_and_trend("GC=F", "Oro")
    oil = get_last_price_and_trend("CL=F", "Petrolio")

    outlook, market_mode, extra_notes = build_market_outlook(
        sp500["trend"],
        nasdaq["trend"],
        gold["trend"],
        oil["trend"]
    )

    profile_data = choose_profile_data(profile.lower(), market_mode)
    strategy = build_strategy_text(profile_data, market_mode)
    news = get_top_news()
    analyses, fractal_summary, traffic_summary = build_fractal_section()
    operational_actions = build_operational_actions(market_mode, analyses)
    macro_table = build_macro_table()

    previous_snapshot = get_latest_snapshot(profile)

    base_allocation_changed = build_allocation_change_flags(
        profile_data["base_allocation"],
        previous_snapshot.get("base_allocation") if previous_snapshot else None
    )
    final_allocation_changed = build_allocation_change_flags(
        profile_data["final_allocation"],
        previous_snapshot.get("final_allocation") if previous_snapshot else None
    )

    equity_changes = build_list_change_flags(
        profile_data["equity_funds"],
        previous_snapshot.get("equity_funds") if previous_snapshot else None
    )
    bond_changes = build_list_change_flags(
        profile_data["bond_funds"],
        previous_snapshot.get("bond_funds") if previous_snapshot else None
    )
    commodity_changes = build_list_change_flags(
        profile_data["commodity_funds"],
        previous_snapshot.get("commodity_funds") if previous_snapshot else None
    )
    abs_return_changes = build_list_change_flags(
        profile_data["abs_return_funds"],
        previous_snapshot.get("abs_return_funds") if previous_snapshot else None
    )

    outlook_changed = build_outlook_change_flag(
        outlook,
        previous_snapshot.get("outlook") if previous_snapshot else None
    )

    fractal_changed_flags = build_fractal_change_flags(
        analyses,
        previous_snapshot.get("fractal_lights") if previous_snapshot else None
    )

    return {
        "timestamp": datetime.now().strftime("%d/%m/%Y %H:%M:%S"),
        "profile_label": profile_data["label"],
        "markets": [sp500, nasdaq, gold, oil],
        "outlook": outlook,
        "outlook_changed": outlook_changed,
        "strategy": strategy,
        "equity_funds": profile_data["equity_funds"],
        "bond_funds": profile_data["bond_funds"],
        "commodity_funds": profile_data["commodity_funds"],
        "abs_return_funds": profile_data["abs_return_funds"],
        "equity_changes": equity_changes,
        "bond_changes": bond_changes,
        "commodity_changes": commodity_changes,
        "abs_return_changes": abs_return_changes,
        "base_allocation": profile_data["base_allocation"],
        "final_allocation": profile_data["final_allocation"],
        "base_allocation_changed": base_allocation_changed,
        "final_allocation_changed": final_allocation_changed,
        "intermarket_notes": extra_notes,
        "fractal_summary": fractal_summary,
        "traffic_summary": traffic_summary,
        "fractal_analyses": analyses,
        "fractal_changed_flags": fractal_changed_flags,
        "operational_actions": operational_actions,
        "macro_table": macro_table,
        "news": news,
        "has_previous_snapshot": previous_snapshot is not None,
    }


def build_change_section(data):
    lines = []

    if data["outlook_changed"]:
        lines.append("- Outlook cambiato *")

    mapping = [
        ("Azionari", data["equity_changes"]),
        ("Obbligazionari", data["bond_changes"]),
        ("Commodities", data["commodity_changes"]),
        ("Absolute / Total Return", data["abs_return_changes"]),
    ]

    for label, changes in mapping:
        for item in changes["entered"]:
            lines.append(f"- {label}: entrato {item} *")
        for item in changes["exited"]:
            lines.append(f"- {label}: uscito {item} *")

    for item in data["fractal_analyses"]:
        if data["fractal_changed_flags"].get(item["label"], False):
            lines.append(f"- Semaforo tecnico cambiato su {item['label']} *")

    if not lines:
        lines.append("- Nessuna variazione rilevante rispetto all'ultimo report archiviato.")

    return "\n".join(lines)


def build_macro_text_table(rows):
    lines = []
    for r in rows:
        lines.append(
            f"""{r['country']}
- Inflazione: {r['inflation']}
- GDP: {r['gdp']}
- Tassi: {r['rates']}
- Consumi: {r['consumption']}
- Disoccupazione: {r['unemployment']}
- Posti di lavoro (proxy): {r['jobs']}
- Leading interno: {r['leading']}
- Commento: {r['comment']}"""
        )
    return "\n\n".join(lines)


def allocation_lines(allocation_dict, changed_flags=None):
    lines = []
    for k, v in allocation_dict.items():
        marker = " *" if changed_flags and changed_flags.get(k, False) else ""
        lines.append(f"- {k}: {v}%{marker}")
    return "\n".join(lines)


def build_text_report(data):
    markets_lines = "\n".join(
        [f"- {m['name']}: {m['last']} ({m['trend']})" for m in data["markets"]]
    )
    notes_text = list_lines(data["intermarket_notes"])
    news_text = list_lines(data["news"])
    actions_text = list_lines(data["operational_actions"])
    equity_text = list_lines(data["equity_funds"])
    bond_text = list_lines(data["bond_funds"])
    commodity_text = list_lines(data["commodity_funds"])
    abs_return_text = list_lines(data["abs_return_funds"])
    changes_text = build_change_section(data)
    macro_text = build_macro_text_table(data["macro_table"])

    fractal_lines = []
    for item in data["fractal_analyses"]:
        marker = " *" if data["fractal_changed_flags"].get(item["label"], False) else ""
        fractal_lines.append(
            f"""- {item['light']} {item['label']} ({item['light_text']}){marker}
  Prezzo: {item['price']}
  Frattale rialzista: {item['up_fractal']}
  Frattale ribassista: {item['down_fractal']}
  Stato: {item['status']}
  Nota: {item['comment']}"""
        )
    fractal_text = "\n".join(fractal_lines)

    note_compare = (
        "Nota confronto: l'asterisco * evidenzia le voci cambiate rispetto all'ultimo report archiviato dello stesso profilo."
        if data["has_previous_snapshot"]
        else "Nota confronto: non esiste ancora un report precedente da confrontare per questo profilo."
    )

    return f"""
REPORT CONSULENZA - {data['timestamp']}

PROFILO CLIENTE:
{data['profile_label']}

VARIAZIONI RISPETTO ALL'ULTIMO REPORT:
{changes_text}

MERCATI:
{markets_lines}

OUTLOOK:
{data['outlook']}{' *' if data['outlook_changed'] else ''}

STRATEGIA:
{data['strategy']}

STRUMENTI AZIONARI SUGGERITI:
{equity_text}

STRUMENTI OBBLIGAZIONARI SUGGERITI:
{bond_text}

STRUMENTI COMMODITIES SUGGERITI:
{commodity_text}

STRUMENTI ABSOLUTE / TOTAL RETURN:
{abs_return_text}

ASSET ALLOCATION BASE:
{allocation_lines(data['base_allocation'], data['base_allocation_changed'])}

ASSET ALLOCATION TATTICA:
{allocation_lines(data['final_allocation'], data['final_allocation_changed'])}

MACROECONOMIA E PROIEZIONI:
{macro_text}

NOTE INTERMARKET:
{notes_text}

ANALISI TECNICA / FRATTALI:
{data['fractal_summary']}
{data['traffic_summary']}

{fractal_text}

AZIONI OPERATIVE:
{actions_text}

NEWS RILEVANTI:
{news_text}

{note_compare}

{DISCLAIMER_TEXT}
"""


def save_text_report(profile: str, report_text: str, prefix=""):
    now = datetime.now()
    safe_profile = profile.lower().strip().replace(" ", "_")
    day_folder = ARCHIVE_DIR / now.strftime("%Y-%m-%d")
    day_folder.mkdir(parents=True, exist_ok=True)

    base_name = f"{now.strftime('%H-%M-%S')}_{prefix}{safe_profile}".replace("__", "_")
    file_path = day_folder / f"{base_name}.txt"

    with open(file_path, "w", encoding="utf-8") as f:
        f.write(report_text)

    return file_path


def save_docx_report(txt_path: Path, report_text: str):
    docx_path = txt_path.with_suffix(".docx")
    doc = Document()

    if LOGO_PATH.exists():
        try:
            doc.add_picture(str(LOGO_PATH), width=Inches(2.0))
        except Exception:
            pass

    doc.add_heading("REPORT CONSULENZA", level=1)
    doc.add_paragraph(f"Data e ora: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}")

    for block in report_text.strip().split("\n\n"):
        doc.add_paragraph(block)

    doc.save(docx_path)
    return docx_path


def section_title(text, styles):
    return Paragraph(text, styles["section"])


def normal_par(text, styles):
    return Paragraph(clean_text_for_pdf(text).replace("\n", "<br/>"), styles["body"])


def bullet_par(text, styles):
    return Paragraph(f"• {clean_text_for_pdf(text)}", styles["body"])


def build_macro_pdf_table(rows):
    table_rows = [[
        "Paese", "Infl.", "GDP", "Tassi", "Consumi", "Disocc.", "Jobs", "Leading", "Commento"
    ]]

    for r in rows:
        table_rows.append([
            r["country"],
            clean_text_for_pdf(r["inflation"]),
            clean_text_for_pdf(r["gdp"]),
            clean_text_for_pdf(r["rates"]),
            clean_text_for_pdf(r["consumption"]),
            clean_text_for_pdf(r["unemployment"]),
            clean_text_for_pdf(r["jobs"]),
            clean_text_for_pdf(r["leading"]),
            clean_text_for_pdf(r["comment"]),
        ])

    return table_rows


def allocation_table_rows(base_alloc, final_alloc, base_flags, final_flags):
    rows = [["Categoria", "Base", "Tattica"]]
    categories = list(dict.fromkeys(list(base_alloc.keys()) + list(final_alloc.keys())))

    for cat in categories:
        base_val = f"{base_alloc.get(cat, '')}%"
        final_val = f"{final_alloc.get(cat, '')}%"

        if base_flags.get(cat, False):
            base_val = f"{base_val} *"
        if final_flags.get(cat, False):
            final_val = f"{final_val} *"

        rows.append([cat, base_val, final_val])

    return rows


def _pdf_styles():
    styles = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "title",
            parent=styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=19,
            textColor=colors.HexColor("#0f3d91"),
            spaceAfter=8,
        ),
        "subtitle": ParagraphStyle(
            "subtitle",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            textColor=colors.HexColor("#5f6b7a"),
            spaceAfter=10,
        ),
        "section": ParagraphStyle(
            "section",
            parent=styles["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=11.5,
            textColor=colors.white,
            backColor=colors.HexColor("#0f3d91"),
            spaceBefore=8,
            spaceAfter=6,
            leftIndent=4,
        ),
        "body": ParagraphStyle(
            "body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=8.8,
            leading=12,
            textColor=colors.HexColor("#1f2937"),
            spaceAfter=4,
        ),
        "small": ParagraphStyle(
            "small",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=8.0,
            leading=10.5,
            textColor=colors.HexColor("#6b7280"),
            spaceBefore=10,
        ),
    }


def save_pdf_report(txt_path: Path, data: dict):
    pdf_path = txt_path.with_suffix(".pdf")
    styles = _pdf_styles()

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm
    )

    story = []

    if LOGO_PATH.exists():
        try:
            logo = Image(str(LOGO_PATH))
            logo.drawHeight = 18 * mm
            logo.drawWidth = 42 * mm
            story.append(logo)
            story.append(Spacer(1, 5))
        except Exception:
            pass

    story.append(Paragraph("REPORT CONSULENZA", styles["title"]))
    story.append(Paragraph(f"Data e ora: {data['timestamp']}", styles["subtitle"]))

    profile_table = Table(
        [["Profilo cliente", data["profile_label"]]],
        colWidths=[42 * mm, 128 * mm]
    )
    profile_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, 0), colors.HexColor("#0f3d91")),
        ("TEXTCOLOR", (0, 0), (0, 0), colors.white),
        ("BACKGROUND", (1, 0), (1, 0), colors.HexColor("#eef4ff")),
        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#b7c8ea")),
        ("INNERGRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#b7c8ea")),
        ("FONTNAME", (0, 0), (0, 0), "Helvetica-Bold"),
        ("FONTNAME", (1, 0), (1, 0), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("PADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(profile_table)
    story.append(Spacer(1, 8))

    story.append(section_title("Variazioni rispetto all'ultimo report", styles))
    for line in build_change_section(data).split("\n"):
        story.append(bullet_par(line.replace("- ", "", 1), styles))

    story.append(section_title("Mercati", styles))
    market_rows = [["Strumento", "Valore", "Trend"]]
    for m in data["markets"]:
        market_rows.append([m["name"], str(m["last"]), m["trend"]])

    markets_table = Table(market_rows, colWidths=[65 * mm, 38 * mm, 40 * mm])
    markets_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d9e8ff")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f3d91")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#c8d6ec")),
        ("PADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(markets_table)

    story.append(section_title("Outlook", styles))
    outlook_text = data["outlook"] + (" *" if data["outlook_changed"] else "")
    story.append(normal_par(outlook_text, styles))

    story.append(section_title("Strategia", styles))
    story.append(normal_par(data["strategy"], styles))

    story.append(section_title("Strumenti azionari suggeriti", styles))
    for item in data["equity_funds"]:
        marker = " *" if item in data["equity_changes"]["entered"] else ""
        story.append(bullet_par(item + marker, styles))
    for item in data["equity_changes"]["exited"]:
        story.append(bullet_par(f"Uscito: {item} *", styles))

    story.append(section_title("Strumenti obbligazionari suggeriti", styles))
    for item in data["bond_funds"]:
        marker = " *" if item in data["bond_changes"]["entered"] else ""
        story.append(bullet_par(item + marker, styles))
    for item in data["bond_changes"]["exited"]:
        story.append(bullet_par(f"Uscito: {item} *", styles))

    story.append(section_title("Strumenti commodities suggeriti", styles))
    for item in data["commodity_funds"]:
        marker = " *" if item in data["commodity_changes"]["entered"] else ""
        story.append(bullet_par(item + marker, styles))
    for item in data["commodity_changes"]["exited"]:
        story.append(bullet_par(f"Uscito: {item} *", styles))

    if data["abs_return_funds"]:
        story.append(section_title("Absolute / Total Return", styles))
        for item in data["abs_return_funds"]:
            marker = " *" if item in data["abs_return_changes"]["entered"] else ""
            story.append(bullet_par(item + marker, styles))

    for item in data["abs_return_changes"]["exited"]:
        story.append(bullet_par(f"Uscito: {item} *", styles))

    story.append(section_title("Asset allocation", styles))
    alloc_rows = allocation_table_rows(
        data["base_allocation"],
        data["final_allocation"],
        data["base_allocation_changed"],
        data["final_allocation_changed"]
    )
    alloc_table = Table(alloc_rows, colWidths=[68 * mm, 38 * mm, 48 * mm])
    alloc_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#fff0cc")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#8a5a00")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#e4d6a7")),
        ("PADDING", (0, 0), (-1, -1), 5),
    ]))
    story.append(alloc_table)

    story.append(section_title("Macroeconomia e proiezioni", styles))
    macro_rows = build_macro_pdf_table(data["macro_table"])
    macro_table = Table(
        macro_rows,
        colWidths=[18*mm, 16*mm, 16*mm, 16*mm, 18*mm, 18*mm, 18*mm, 16*mm, 52*mm],
        repeatRows=1
    )
    macro_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dff6e5")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f5132")),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "Helvetica"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.4),
        ("LEADING", (0, 0), (-1, -1), 7.0),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#b8d9c3")),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("LEFTPADDING", (0, 0), (-1, -1), 3),
        ("RIGHTPADDING", (0, 0), (-1, -1), 3),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(macro_table)

    story.append(section_title("Note intermarket", styles))
    for note in data["intermarket_notes"]:
        story.append(bullet_par(note, styles))

    story.append(section_title("Analisi tecnica / frattali", styles))
    story.append(normal_par(data["fractal_summary"], styles))
    story.append(normal_par(data["traffic_summary"], styles))

    for item in data["fractal_analyses"]:
        marker = " *" if data["fractal_changed_flags"].get(item["label"], False) else ""
        story.append(normal_par(
            f"{item['label']} - {item['light_text']}{marker}\n"
            f"Prezzo: {item['price']}\n"
            f"Frattale rialzista: {item['up_fractal']}\n"
            f"Frattale ribassista: {item['down_fractal']}\n"
            f"Stato: {item['status']}\n"
            f"Nota: {item['comment']}",
            styles
        ))

    story.append(section_title("Azioni operative", styles))
    for action in data["operational_actions"]:
        story.append(bullet_par(action, styles))

    story.append(section_title("News rilevanti", styles))
    for news in data["news"]:
        story.append(bullet_par(news, styles))

    story.append(Spacer(1, 8))
    story.append(Paragraph(clean_text_for_pdf(DISCLAIMER_TEXT), styles["small"]))

    doc.build(story)
    return pdf_path


def create_pie_chart(allocation_dict, title, output_path: Path):
    labels = list(allocation_dict.keys())
    values = list(allocation_dict.values())

    plt.figure(figsize=(4.2, 4.2))
    plt.pie(values, labels=labels, autopct="%1.0f%%", startangle=90)
    plt.title(title)
    plt.tight_layout()
    plt.savefig(output_path, dpi=150)
    plt.close()


def save_master_profiles_pdf():
    now = datetime.now()
    day_folder = ARCHIVE_DIR / now.strftime("%Y-%m-%d")
    day_folder.mkdir(parents=True, exist_ok=True)

    pdf_path = day_folder / f"{now.strftime('%H-%M-%S')}_report_3_profili.pdf"
    styles = _pdf_styles()

    profiles_data = {
        key: build_dashboard_data(key)
        for key in ["prudente", "bilanciato", "dinamico"]
    }

    chart_paths = []
    for key, data in profiles_data.items():
        chart_path = day_folder / f"{now.strftime('%H-%M-%S')}_{key}_pie.png"
        create_pie_chart(data["final_allocation"], f"{data['profile_label']} - Allocation tattica", chart_path)
        chart_paths.append(chart_path)

    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        leftMargin=15 * mm,
        rightMargin=15 * mm,
        topMargin=15 * mm,
        bottomMargin=15 * mm
    )

    story = []

    if LOGO_PATH.exists():
        try:
            logo = Image(str(LOGO_PATH))
            logo.drawHeight = 18 * mm
            logo.drawWidth = 42 * mm
            story.append(logo)
            story.append(Spacer(1, 5))
        except Exception:
            pass

    story.append(Paragraph("REPORT UNICO DEI 3 PROFILI", styles["title"]))
    story.append(Paragraph(f"Data e ora: {now.strftime('%d/%m/%Y %H:%M:%S')}", styles["subtitle"]))

    for idx, key in enumerate(["prudente", "bilanciato", "dinamico"]):
        data = profiles_data[key]
        chart_path = chart_paths[idx]

        story.append(section_title(data["profile_label"], styles))
        story.append(normal_par(data["outlook"], styles))
        story.append(normal_par(data["strategy"], styles))

        alloc_table = Table(
            [["Categoria", "Peso"]] + [[k, f"{v}%"] for k, v in data["final_allocation"].items()],
            colWidths=[70 * mm, 35 * mm]
        )
        alloc_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef4ff")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f3d91")),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#c8d6ec")),
            ("PADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(alloc_table)
        story.append(Spacer(1, 6))

        try:
            img = Image(str(chart_path))
            img.drawWidth = 85 * mm
            img.drawHeight = 85 * mm
            story.append(img)
        except Exception:
            story.append(normal_par("Grafico a torta non disponibile.", styles))

        story.append(Spacer(1, 6))
        story.append(normal_par("Strumenti azionari:", styles))
        for item in data["equity_funds"]:
            story.append(bullet_par(item, styles))

        story.append(normal_par("Strumenti obbligazionari:", styles))
        for item in data["bond_funds"]:
            story.append(bullet_par(item, styles))

        story.append(normal_par("Commodities:", styles))
        for item in data["commodity_funds"]:
            story.append(bullet_par(item, styles))

        if data["abs_return_funds"]:
            story.append(normal_par("Absolute / Total Return:", styles))
            for item in data["abs_return_funds"]:
                story.append(bullet_par(item, styles))

        if idx < 2:
            story.append(PageBreak())

    story.append(PageBreak())
    story.append(section_title("Disclaimer", styles))
    story.append(normal_par(DISCLAIMER_TEXT, styles))

    doc.build(story)
    return pdf_path


def list_archived_reports():
    files = []
    for path in sorted(ARCHIVE_DIR.rglob("*")):
        if path.is_file() and path.suffix.lower() in [".txt", ".docx", ".pdf", ".png", ".json"]:
            files.append({
                "name": path.name,
                "date_folder": path.parent.name,
                "relative_path": str(path.relative_to(ARCHIVE_DIR)),
                "extension": path.suffix.lower()
            })
    files.reverse()
    return [f for f in files if f["extension"] not in [".json", ".png"]][:300]


@app.get("/", response_class=HTMLResponse)
def home():
    return """
    <!DOCTYPE html>
    <html lang="it">
    <head>
        <meta charset="utf-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>CONSULENTIA AI</title>
        <style>
            :root{
                --bg:#f4f7fb; --card:#ffffff; --text:#1f2937; --muted:#6b7280;
                --border:#e5e7eb; --shadow:0 6px 20px rgba(0,0,0,0.06);
                --blue:#2563eb; --green:#16a34a; --yellow:#ca8a04; --red:#dc2626; --slate:#475569;
                --orange:#d97706;
            }
            body{font-family:Arial,sans-serif;margin:0;background:var(--bg);color:var(--text);}
            .wrap{max-width:1240px;margin:0 auto;padding:20px;}
            .topbar{display:flex;flex-wrap:wrap;gap:12px;align-items:center;justify-content:space-between;margin-bottom:20px;}
            .title{font-size:32px;font-weight:700;}
            .controls{display:flex;gap:10px;flex-wrap:wrap;align-items:center;}
            select,button{padding:12px 14px;border-radius:12px;border:1px solid var(--border);font-size:16px;}
            button{cursor:pointer;background:#111827;color:white;border:none;}
            button.secondary{background:#4b5563;}
            .status{margin:8px 0 18px;color:var(--muted);font-size:14px;}
            .grid{display:grid;grid-template-columns:repeat(auto-fit, minmax(320px, 1fr));gap:16px;}
            .card{background:var(--card);border-radius:18px;padding:18px;box-shadow:var(--shadow);border:1px solid var(--border);}
            .card h3{margin:0 0 12px 0;font-size:18px;}
            .full{grid-column:1 / -1;}
            .card.blue{border-left:6px solid var(--blue);}
            .card.green{border-left:6px solid var(--green);}
            .card.yellow{border-left:6px solid var(--yellow);}
            .card.red{border-left:6px solid var(--red);}
            .card.slate{border-left:6px solid var(--slate);}
            .market-item,.list-item,.alloc-item{padding:10px 0;border-bottom:1px solid #eef2f7;}
            .market-item:last-child,.list-item:last-child,.alloc-item:last-child{border-bottom:none;}
            .label{font-weight:700;}
            .small{color:var(--muted);font-size:14px;}
            .pill{display:inline-block;padding:5px 10px;border-radius:999px;background:#eef2ff;font-size:13px;margin-top:6px;}
            .mono{white-space:pre-wrap;line-height:1.45;}
            .market-value.pos{color:var(--green);font-weight:700;}
            .market-value.neg{color:var(--red);font-weight:700;}
            .signal{display:inline-block;padding:6px 10px;border-radius:999px;font-size:13px;font-weight:700;margin-bottom:8px;}
            .signal.green{background:#dcfce7;color:#166534;}
            .signal.yellow{background:#fef9c3;color:#854d0e;}
            .signal.red{background:#fee2e2;color:#991b1b;}
            .signal.gray{background:#e5e7eb;color:#374151;}
            .changed{color:var(--orange);font-weight:700;}
            .changed-star{color:var(--orange);font-weight:700;margin-left:8px;}
            .note-change{margin-top:8px;font-size:13px;color:var(--orange);font-weight:700;}
            .macro-wrap{overflow-x:auto;}
            .macro-table{width:100%; border-collapse:collapse; font-size:13px;}
            .macro-table th,.macro-table td{border:1px solid #e5e7eb; padding:8px; vertical-align:top; text-align:left;}
            .macro-table th{background:#eef4ff;}
            #textReportBox{display:none;margin-top:18px;background:#111827;color:#f9fafb;padding:18px;border-radius:16px;white-space:pre-wrap;line-height:1.45;}
            .archive-item{padding:10px 0;border-bottom:1px solid #eef2f7;font-size:14px;}
        </style>
    </head>
    <body>
        <div class="wrap">
            <div class="topbar">
                <div>
                    <div class="title">CONSULENTIA AI</div>
                    <div class="small">Dashboard consulenza con confronto, macro, tecnica e report unico 3 profili</div>
                </div>
                <div class="controls">
                    <label for="profile"><strong>Profilo cliente</strong></label>
                    <select id="profile">
                        <option value="prudente">Prudente</option>
                        <option value="bilanciato" selected>Bilanciato</option>
                        <option value="dinamico">Dinamico</option>
                    </select>
                    <button onclick="loadDashboard()">Genera Dashboard</button>
                    <button class="secondary" onclick="loadTextReport()">Mostra report testuale completo</button>
                    <button class="secondary" onclick="exportReport('txt')">Esporta TXT</button>
                    <button class="secondary" onclick="exportReport('docx')">Esporta Word</button>
                    <button class="secondary" onclick="exportReport('pdf')">Esporta PDF</button>
                    <button class="secondary" onclick="exportMasterPdf()">Report 3 profili PDF</button>
                    <button class="secondary" onclick="loadArchive()">Aggiorna archivio</button>
                </div>
            </div>

            <div id="status" class="status">Seleziona il profilo e premi "Genera Dashboard".</div>
            <div id="dashboard" class="grid"></div>
            <div id="textReportBox"></div>

            <div class="card full slate" style="margin-top:18px;">
                <h3>Archivio report</h3>
                <div id="archiveBox" class="small">Nessun archivio caricato.</div>
            </div>
        </div>

        <script>
            function escapeHtml(text) {
                if (text === null || text === undefined) return "";
                return String(text)
                    .replaceAll("&", "&amp;")
                    .replaceAll("<", "&lt;")
                    .replaceAll(">", "&gt;");
            }

            function trendClass(trend) {
                if (trend === "positivo") return "pos";
                if (trend === "negativo") return "neg";
                return "";
            }

            function signalClass(light) {
                if (light === "🟢") return "green";
                if (light === "🟡") return "yellow";
                if (light === "🔴") return "red";
                return "gray";
            }

            function renderList(items) {
                return items.map(item => `<div class="list-item">${escapeHtml(item)}</div>`).join("");
            }

            function renderChangedList(currentItems, changes) {
                const entered = new Set((changes && changes.entered) || []);
                const exited = (changes && changes.exited) || [];
                let html = "";

                currentItems.forEach(item => {
                    html += `
                        <div class="list-item">
                            ${escapeHtml(item)}
                            ${entered.has(item) ? '<span class="changed-star">*</span>' : ''}
                        </div>
                    `;
                });

                exited.forEach(item => {
                    html += `
                        <div class="list-item changed">
                            Uscito: ${escapeHtml(item)} <span class="changed-star">*</span>
                        </div>
                    `;
                });

                return html || '<div class="list-item">Nessuno</div>';
            }

            function renderMarkets(markets) {
                return markets.map(m => `
                    <div class="market-item">
                        <div class="label">${escapeHtml(m.name)}</div>
                        <div class="market-value ${trendClass(m.trend)}">${escapeHtml(m.last)} <span class="small">(${escapeHtml(m.trend)})</span></div>
                    </div>
                `).join("");
            }

            function renderFractals(items, changedFlags) {
                return items.map(item => `
                    <div class="list-item">
                        <div class="signal ${signalClass(item.light)}">${escapeHtml(item.light)} ${escapeHtml(item.light_text)}</div>
                        <div class="label">
                            ${escapeHtml(item.label)}
                            ${changedFlags && changedFlags[item.label] ? '<span class="changed-star">*</span>' : ''}
                        </div>
                        <div class="small">Prezzo: ${escapeHtml(item.price)}</div>
                        <div class="small">Frattale rialzista: ${escapeHtml(item.up_fractal)}</div>
                        <div class="small">Frattale ribassista: ${escapeHtml(item.down_fractal)}</div>
                        <div class="small">Stato: ${escapeHtml(item.status)}</div>
                        <div>${escapeHtml(item.comment)}</div>
                    </div>
                `).join("");
            }

            function renderAllocation(obj, changedFlags) {
                return Object.entries(obj).map(([k,v]) => `
                    <div class="alloc-item">
                        <strong>${escapeHtml(k)}</strong>:
                        <span class="${changedFlags && changedFlags[k] ? 'changed' : ''}">${escapeHtml(v)}%</span>
                        ${changedFlags && changedFlags[k] ? '<span class="changed-star">*</span>' : ''}
                    </div>
                `).join("");
            }

            function renderVariationBox(data) {
                const lines = [];

                if (data.outlook_changed) lines.push('Outlook cambiato *');

                [
                    ['Azionari', data.equity_changes],
                    ['Obbligazionari', data.bond_changes],
                    ['Commodities', data.commodity_changes],
                    ['Absolute / Total Return', data.abs_return_changes]
                ].forEach(([label, ch]) => {
                    (ch.entered || []).forEach(x => lines.push(`${label}: entrato ${x} *`));
                    (ch.exited || []).forEach(x => lines.push(`${label}: uscito ${x} *`));
                });

                (data.fractal_analyses || []).forEach(item => {
                    if (data.fractal_changed_flags && data.fractal_changed_flags[item.label]) {
                        lines.push(`Semaforo tecnico cambiato su ${item.label} *`);
                    }
                });

                if (!lines.length) {
                    return '<div class="list-item">Nessuna variazione rilevante rispetto all\\'ultimo report archiviato.</div>';
                }

                return lines.map(x => `<div class="list-item changed">${escapeHtml(x)}</div>`).join('');
            }

            function renderMacroTable(rows) {
                if (!rows || !rows.length) {
                    return '<div class="list-item">Nessun dato macro disponibile.</div>';
                }

                let html = `
                    <div class="macro-wrap">
                        <table class="macro-table">
                            <thead>
                                <tr>
                                    <th>Paese</th>
                                    <th>Inflazione</th>
                                    <th>GDP</th>
                                    <th>Tassi</th>
                                    <th>Consumi</th>
                                    <th>Disoccupazione</th>
                                    <th>Posti di lavoro</th>
                                    <th>Leading</th>
                                    <th>Commento</th>
                                </tr>
                            </thead>
                            <tbody>
                `;

                rows.forEach(r => {
                    html += `
                        <tr>
                            <td><strong>${escapeHtml(r.country)}</strong></td>
                            <td>${escapeHtml(r.inflation)}</td>
                            <td>${escapeHtml(r.gdp)}</td>
                            <td>${escapeHtml(r.rates)}</td>
                            <td>${escapeHtml(r.consumption)}</td>
                            <td>${escapeHtml(r.unemployment)}</td>
                            <td>${escapeHtml(r.jobs)}</td>
                            <td>${escapeHtml(r.leading)}</td>
                            <td>${escapeHtml(r.comment)}</td>
                        </tr>
                    `;
                });

                html += "</tbody></table></div>";
                return html;
            }

            function renderDashboard(data) {
                const changeNote = data.has_previous_snapshot
                    ? '<div class="note-change">* evidenzia le voci cambiate rispetto all\\'ultimo report archiviato dello stesso profilo.</div>'
                    : '<div class="small" style="margin-top:8px;">Non esiste ancora un report precedente da confrontare per questo profilo.</div>';

                const html = `
                    <div class="card full yellow">
                        <h3>Variazioni rispetto all'ultimo report</h3>
                        ${renderVariationBox(data)}
                        ${changeNote}
                    </div>

                    <div class="card blue">
                        <h3>Profilo</h3>
                        <div class="label">${escapeHtml(data.profile_label)}</div>
                        <div class="pill">${escapeHtml(data.timestamp)}</div>
                    </div>

                    <div class="card ${data.outlook.toLowerCase().includes("costruttivo") ? "green" : (data.outlook.toLowerCase().includes("fragile") ? "red" : "yellow")}">
                        <h3>Outlook</h3>
                        <div class="mono ${data.outlook_changed ? 'changed' : ''}">
                            ${escapeHtml(data.outlook)}
                            ${data.outlook_changed ? '<span class="changed-star">*</span>' : ''}
                        </div>
                    </div>

                    <div class="card full blue">
                        <h3>Mercati</h3>
                        ${renderMarkets(data.markets)}
                    </div>

                    <div class="card full slate">
                        <h3>Strategia</h3>
                        <div class="mono">${escapeHtml(data.strategy)}</div>
                    </div>

                    <div class="card full green">
                        <h3>Strumenti azionari suggeriti</h3>
                        ${renderChangedList(data.equity_funds, data.equity_changes)}
                    </div>

                    <div class="card full blue">
                        <h3>Strumenti obbligazionari suggeriti</h3>
                        ${renderChangedList(data.bond_funds, data.bond_changes)}
                    </div>

                    <div class="card full yellow">
                        <h3>Strumenti commodities suggeriti</h3>
                        ${renderChangedList(data.commodity_funds, data.commodity_changes)}
                    </div>

                    <div class="card full slate">
                        <h3>Absolute / Total Return</h3>
                        ${renderChangedList(data.abs_return_funds && data.abs_return_funds.length ? data.abs_return_funds : ["Nessuno previsto per questo profilo"], data.abs_return_changes)}
                    </div>

                    <div class="card yellow">
                        <h3>Allocation base</h3>
                        ${renderAllocation(data.base_allocation, data.base_allocation_changed)}
                    </div>

                    <div class="card ${data.profile_label === "Dinamico" ? "green" : "slate"}">
                        <h3>Allocation tattica</h3>
                        ${renderAllocation(data.final_allocation, data.final_allocation_changed)}
                    </div>

                    <div class="card full green">
                        <h3>Macroeconomia e proiezioni</h3>
                        ${renderMacroTable(data.macro_table)}
                    </div>

                    <div class="card full yellow">
                        <h3>Note intermarket</h3>
                        ${renderList(data.intermarket_notes)}
                    </div>

                    <div class="card full ${data.fractal_summary.toLowerCase().includes("favorevole") ? "green" : (data.fractal_summary.toLowerCase().includes("debole") ? "red" : "yellow")}">
                        <h3>Analisi tecnica / frattali</h3>
                        <div class="label">${escapeHtml(data.fractal_summary)}</div>
                        <div class="pill">${escapeHtml(data.traffic_summary)}</div>
                        <div style="margin-top:12px;">
                            ${renderFractals(data.fractal_analyses, data.fractal_changed_flags)}
                        </div>
                    </div>

                    <div class="card full green">
                        <h3>Azioni operative</h3>
                        ${renderList(data.operational_actions)}
                    </div>

                    <div class="card full red">
                        <h3>News rilevanti</h3>
                        ${renderList(data.news)}
                    </div>
                `;
                document.getElementById("dashboard").innerHTML = html;
            }

            function renderArchive(items) {
                if (!items || items.length === 0) {
                    document.getElementById("archiveBox").innerHTML = "Archivio vuoto.";
                    return;
                }

                const html = items.map(item => `
                    <div class="archive-item">
                        <div><strong>${escapeHtml(item.name)}</strong></div>
                        <div class="small">${escapeHtml(item.date_folder)} — ${escapeHtml(item.relative_path)}</div>
                    </div>
                `).join("");

                document.getElementById("archiveBox").innerHTML = html;
            }

            async function loadDashboard() {
                const profile = document.getElementById("profile").value;
                document.getElementById("status").innerText = "Generazione dashboard in corso...";
                document.getElementById("textReportBox").style.display = "none";

                try {
                    const res = await fetch("/api/dashboard?profile=" + encodeURIComponent(profile));
                    const data = await res.json();
                    renderDashboard(data);
                    document.getElementById("status").innerText = "Dashboard aggiornata.";
                } catch (err) {
                    document.getElementById("status").innerText = "Errore: " + err;
                }
            }

            async function loadTextReport() {
                const profile = document.getElementById("profile").value;
                document.getElementById("status").innerText = "Generazione report testuale in corso...";

                try {
                    const res = await fetch("/report?profile=" + encodeURIComponent(profile));
                    const data = await res.json();
                    const box = document.getElementById("textReportBox");
                    box.innerText = data.report;
                    box.style.display = "block";
                    document.getElementById("status").innerText = "Report testuale pronto e archiviato.";
                    window.scrollTo({ top: document.body.scrollHeight, behavior: "smooth" });
                    loadArchive();
                } catch (err) {
                    document.getElementById("status").innerText = "Errore: " + err;
                }
            }

            async function exportReport(format) {
                const profile = document.getElementById("profile").value;
                document.getElementById("status").innerText = "Esportazione in corso...";

                try {
                    const res = await fetch("/export?profile=" + encodeURIComponent(profile) + "&format=" + encodeURIComponent(format));
                    const data = await res.json();

                    if (data.download_url) {
                        window.open(data.download_url, "_blank");
                    }

                    document.getElementById("status").innerText = "Esportazione completata.";
                    loadArchive();
                } catch (err) {
                    document.getElementById("status").innerText = "Errore: " + err;
                }
            }

            async function exportMasterPdf() {
                document.getElementById("status").innerText = "Generazione report unico dei 3 profili in corso...";

                try {
                    const res = await fetch("/export_master_pdf");
                    const data = await res.json();
                    if (data.download_url) {
                        window.open(data.download_url, "_blank");
                    }
                    document.getElementById("status").innerText = "Report unico dei 3 profili pronto.";
                    loadArchive();
                } catch (err) {
                    document.getElementById("status").innerText = "Errore: " + err;
                }
            }

            async function loadArchive() {
                try {
                    const res = await fetch("/api/archive");
                    const data = await res.json();
                    renderArchive(data.items || []);
                } catch (err) {
                    document.getElementById("archiveBox").innerHTML = "Errore caricamento archivio: " + escapeHtml(err);
                }
            }

            loadArchive();
        </script>
    </body>
    </html>
    """


@app.get("/api/dashboard")
def api_dashboard(profile: str = Query(default="bilanciato")):
    try:
        data = build_dashboard_data(profile)
        return JSONResponse(content=data)
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app.get("/api/archive")
def api_archive():
    try:
        items = list_archived_reports()
        return JSONResponse(content={"items": items})
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app.get("/report")
def report(profile: str = Query(default="bilanciato")):
    try:
        data = build_dashboard_data(profile)
        report_text = build_text_report(data)
        txt_path = save_text_report(profile, report_text)
        save_snapshot(profile, data, txt_path)
        return {"report": report_text}
    except Exception as e:
        return {"report": f"Errore nella generazione del report: {e}"}


@app.get("/export")
def export_report(
    profile: str = Query(default="bilanciato"),
    format: str = Query(default="txt")
):
    try:
        data = build_dashboard_data(profile)
        report_text = build_text_report(data)
        txt_path = save_text_report(profile, report_text)
        save_snapshot(profile, data, txt_path)

        format = format.lower().strip()

        if format == "txt":
            final_path = txt_path
        elif format == "docx":
            final_path = save_docx_report(txt_path, report_text)
        elif format == "pdf":
            final_path = save_pdf_report(txt_path, data)
        else:
            return JSONResponse(content={"error": "Formato non supportato"}, status_code=400)

        download_url = f"/download/{final_path.relative_to(ARCHIVE_DIR).as_posix()}"
        return JSONResponse(content={
            "message": "Esportazione completata",
            "download_url": download_url
        })
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app.get("/export_master_pdf")
def export_master_pdf():
    try:
        final_path = save_master_profiles_pdf()
        download_url = f"/download/{final_path.relative_to(ARCHIVE_DIR).as_posix()}"
        return JSONResponse(content={
            "message": "Report unico creato",
            "download_url": download_url
        })
    except Exception as e:
        return JSONResponse(content={"error": str(e)}, status_code=500)


@app.get("/download/{file_path:path}")
def download_file(file_path: str):
    full_path = ARCHIVE_DIR / file_path

    if not full_path.exists() or not full_path.is_file():
        return JSONResponse(content={"error": "File non trovato"}, status_code=404)

    return FileResponse(path=str(full_path), filename=full_path.name)